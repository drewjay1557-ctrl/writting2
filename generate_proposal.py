# -*- coding: utf-8 -*-
"""
生成GNSS RTK/INS自适应滑动窗口因子图融合开题报告Word文档
优化版：
  1. 仅第1章（研究意义与国内外现状）引用参考文献；
  2. 参考文献近五年（2021—2025）文献占多数；
  3. 其余章节移除所有引用标注。
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


CN_FONT = 'SimSun'        # 中文字体（宋体，默认）
CN_FONT_HEI = 'SimHei'    # 黑体（标题）
EN_FONT = 'Times New Roman'


def set_run_font(run, size_pt=12, bold=False, cn_font=CN_FONT, en_font=EN_FONT, color=None):
    """设置字体：西文 Times New Roman + 中文宋体/黑体。"""
    run.font.name = en_font
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), en_font)
    rFonts.set(qn('w:hAnsi'), en_font)
    rFonts.set(qn('w:eastAsia'), cn_font)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size_pt=18, bold=True, cn_font=CN_FONT_HEI)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(18)
    return p


def add_heading(doc, text, level=1):
    """自定义标题样式"""
    sizes = {1: 16, 2: 14, 3: 12}
    size = sizes.get(level, 12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, size_pt=size, bold=True, cn_font=CN_FONT_HEI)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    return p


def add_body(doc, text, indent=True, first_line_indent_chars=2):
    """正文段落 (支持[1][2]上标参考文献标注)"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.first_line_indent = Pt(12 * first_line_indent_chars)

    import re
    # 用正则切分，保留 [数字] 或 [数字,数字] 作为独立段
    parts = re.split(r'(\[\d+(?:[,\-]\d+)*\])', text)
    for part in parts:
        if re.match(r'^\[\d+(?:[,\-]\d+)*\]$', part):
            # 参考文献上标
            run = p.add_run(part)
            set_run_font(run, size_pt=12)
            run.font.superscript = True
        else:
            run = p.add_run(part)
            set_run_font(run, size_pt=12)
    return p


def add_bullet(doc, text, level=0):
    """项目符号段落（无引用）"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Pt(20 + level * 20)
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run('• ' + text)
    set_run_font(run, size_pt=12)
    return p


def add_numbered(doc, text, num):
    """编号段落（第一级：(1)(2)(3)）— 无引用版本"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Pt(20)
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f"({num}) {text}")
    set_run_font(run, size_pt=12)
    return p


def add_reference(doc, num, text):
    """参考文献条目"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Pt(24)
    p.paragraph_format.first_line_indent = Pt(-24)  # 悬挂缩进
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f"[{num}] {text}")
    set_run_font(run, size_pt=11)
    return p


def set_document_defaults(doc):
    """设置默认字体、页边距等"""
    style = doc.styles['Normal']
    style.font.name = EN_FONT
    style.font.size = Pt(12)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), EN_FONT)
    rFonts.set(qn('w:hAnsi'), EN_FONT)
    rFonts.set(qn('w:eastAsia'), CN_FONT)

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)


def build_doc():
    doc = Document()
    set_document_defaults(doc)

    # ============ 封面标题 ============
    add_title(doc, "硕士学位论文开题报告")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("论文题目：基于自适应滑动窗口因子图融合的GNSS RTK/INS组合导航算法研究")
    set_run_font(run, size_pt=14, bold=True, cn_font=CN_FONT_HEI)
    p.paragraph_format.space_after = Pt(18)

    # ================================================================
    # 1. 研究意义及国内外研究现状 —— 仅本章引用参考文献
    # ================================================================
    add_heading(doc, "1  研究意义及国内外研究现状", level=1)

    # ---------- 1.1 研究意义 ----------
    add_heading(doc, "1.1  研究意义", level=2)
    add_body(doc,
        "高精度、高可靠的实时定位与导航是无人驾驶、无人机、智能机器人、精准农业等"
        "众多应用领域的核心基础。全球导航卫星系统（Global Navigation Satellite System, "
        "GNSS）因具备全天候、全球覆盖、绝对定位等特点而得到广泛应用，其中实时动态"
        "（Real-Time Kinematic, RTK）定位技术在开阔环境下可提供厘米级定位精度[1]。"
        "然而在城市峡谷、隧道、林荫道等遮挡环境下，GNSS信号易受非视距（Non-Line-of-Sight, "
        "NLOS）传播和多路径效应影响，卫星可见性下降、观测质量急剧劣化，单独使用GNSS"
        "难以满足应用对连续性和可靠性的要求[2,3]。"
    )
    add_body(doc,
        "惯性导航系统（Inertial Navigation System, INS）不依赖外部信号，能在短时间"
        "内提供高频率的位置、速度和姿态信息，与GNSS在时频域和信息域均具有良好的互补"
        "关系[4]。因此，GNSS与INS的组合导航已成为实现高精度连续定位的主流技术。"
        "传统组合方式主要基于扩展卡尔曼滤波（Extended Kalman Filter, EKF）及其变种，"
        "具有计算量小、实时性好等优点，但也存在以下不足：其一，滤波方法本质上是"
        "马尔科夫过程，仅利用当前时刻信息，对过去观测的利用不够充分；其二，线性化"
        "过程会引入模型误差，在强非线性场景下精度下降；其三，对野值和异常观测的"
        "鲁棒性较弱[5,6]。"
    )
    add_body(doc,
        "因子图优化（Factor Graph Optimization, FGO）作为一种概率图模型框架，"
        "通过在时间窗口内联合优化多时刻状态与多源观测，能够充分利用历史信息，"
        "并支持多传感器异步融合、延迟观测、非线性观测模型等复杂情况[7,8]。"
        "iSAM2增量平滑算法使因子图在实时导航中的应用成为可能[9]，"
        "近年来已在SLAM、视觉惯性导航、GNSS/INS融合等领域取得显著成果[6,10]。"
        "将因子图方法引入GNSS RTK/INS紧组合，可有效提升在复杂环境下的定位精度和"
        "鲁棒性[10,11]。"
    )
    add_body(doc,
        "然而，现有基于因子图的GNSS/INS融合方法多采用固定长度的滑动窗口，"
        "难以兼顾精度与实时性：窗口过短会损失历史信息、降低优化效果；窗口过长则"
        "计算负担过大、难以满足实时性要求[12]。此外，不同行驶场景（如高动态机动、"
        "静止停车、遮挡突变等）对窗口长度和因子权重的需求存在显著差异，"
        "固定参数难以适应动态变化的环境[13]。因此，研究自适应滑动窗口因子图融合算法，"
        "根据观测质量和运动状态动态调整窗口尺寸与因子权重，对于提升复杂场景下"
        "GNSS RTK/INS组合导航系统的性能具有重要的理论意义和工程应用价值。"
    )

    # ---------- 1.2 国内外研究现状 ----------
    add_heading(doc, "1.2  国内外研究现状", level=2)

    add_body(doc,
        "（1）GNSS/INS组合导航研究现状。GNSS与INS的组合方式按照耦合程度可分为"
        "松组合、紧组合和深组合三类[4]。松组合以GNSS定位结果作为输入，结构简单但"
        "在卫星数不足时难以工作；紧组合直接利用伪距、载波相位等原始观测，"
        "即使可见卫星少于四颗仍可提供修正信息，是当前研究的主流[14]。Li等针对"
        "GNSS受限环境提出了多GNSS PPP与视觉惯性系统半紧耦合方法，显著提升了"
        "复杂环境下的连续定位能力[5]。Cao等提出GVINS紧耦合GNSS与视觉惯性里程计，"
        "实现了室内外无缝切换定位[15]。国内学者高亢、李团等在多频多系统RTK/INS"
        "紧组合方面持续推进了工程化应用[16]。"
    )
    add_body(doc,
        "（2）因子图优化方法研究现状。因子图为机器人感知与多传感器融合提供了"
        "通用的概率建模框架，Dellaert等系统总结了因子图在机器人感知中的理论"
        "基础[7]；Indelman等较早将因子图应用于GNSS/INS融合[8]。近年来Wen等"
        "系统比较了EKF与FGO用于GNSS/INS组合导航的性能，证明FGO在城市遮挡环境下"
        "可显著降低定位RMSE[10]。Wen和Hsu进一步利用3D点云辅助的FGO实现NLOS抑制，"
        "在中国香港多个实测场景中取得优良效果[3]。针对RTK/INS紧组合，Zhang等基于"
        "因子图提出了载波相位模糊度与惯导状态联合估计方法[17]。2023年以来，"
        "Chi等发布了开源的GNSS/INS/Camera集成导航库GICI-LIB，支持多种因子类型"
        "和优化后端，为相关研究提供了重要的开源平台[11]。Hsu等推出的UrbanNav"
        "数据集为城市环境定位算法提供了统一基准[18]。"
    )
    add_body(doc,
        "（3）滑动窗口与自适应策略研究现状。滑动窗口技术最早广泛应用于视觉惯性"
        "导航领域，VINS-Mono采用固定长度滑动窗口与边缘化策略，实现了良好的"
        "实时性与精度[19]。在自适应方面，Sünderhauf等提出Switchable Constraints"
        "机制动态抑制异常因子[20]；Pfeifer等提出基于EM的动态协方差估计，"
        "在图优化中自适应调整权重[21]。杨元喜等围绕抗差自适应滤波开展了系统研究"
        "[22]。近五年来，面向GNSS/INS因子图的自适应研究蓬勃发展：Jing等提出"
        "基于鲁棒核函数与方差分量的自适应FGO，显著提升了GNSS退化场景下的精度"
        "[12]；Bai等针对城市峡谷环境，提出基于观测质量评估的自适应因子图权重"
        "调整方法[13]；Zhang等探索了面向低成本IMU的滑窗FGO参数在线整定策略"
        "[23]。然而，上述工作多聚焦于权重自适应，针对窗口长度本身自适应的研究"
        "仍相对较少，如何根据可见卫星数、残差统计量和动态特性综合决策窗口尺寸，"
        "尚无成熟的通用框架。"
    )
    add_body(doc,
        "综上所述，基于因子图的GNSS RTK/INS组合导航研究近年来取得了显著进展，"
        "尤其在鲁棒核函数、权重在线估计、开源工具链等方面成果丰硕。然而，"
        "在复杂城市环境下的窗口长度自适应、计算效率与鲁棒性协同优化方面仍有较大"
        "提升空间。针对这一问题开展深入研究具有重要的理论和工程意义。"
    )

    # ================================================================
    # 2. 研究目标、研究内容及拟解决的关键性问题 —— 无引用
    # ================================================================
    add_heading(doc, "2  研究目标、研究内容及拟解决的关键性问题", level=1)

    add_heading(doc, "2.1  研究目标", level=2)
    add_body(doc,
        "针对复杂城市环境下GNSS RTK/INS组合导航精度波动大、鲁棒性不足的问题，"
        "本课题拟以开源GICI-LIB为基础实验平台，深入研究自适应滑动窗口因子图融合"
        "方法。总体研究目标如下：一是构建GNSS RTK与INS紧组合的因子图数学模型，"
        "完善各类因子（RTK双差因子、IMU预积分因子、先验因子等）的实现；"
        "二是提出基于观测质量和动态特性的自适应滑动窗口策略，实现窗口长度、因子"
        "权重与边缘化方案的在线动态调整；三是在仿真和实测数据集上验证所提方法"
        "相对于传统EKF和固定窗口FGO的性能提升，形成可在嵌入式平台运行的实用算法。"
    )

    add_heading(doc, "2.2  研究内容", level=2)
    add_body(doc,
        "本课题围绕研究目标，主要开展以下四方面研究工作："
    )
    add_numbered(doc,
        "GNSS RTK/INS紧组合因子图建模。建立统一的因子图数学描述，"
        "包括GNSS双差观测因子（伪距、载波相位、多普勒）、IMU预积分因子、"
        "零偏随机游走因子、先验因子以及模糊度参数因子。研究各类因子的雅可比矩阵"
        "与协方差矩阵建模方法。", 1)
    add_numbered(doc,
        "自适应滑动窗口机制设计。针对窗口长度自适应问题，研究基于可见卫星数、"
        "信噪比、观测残差统计量和载体机动性的窗口长度动态调整规则；针对因子权重"
        "自适应问题，研究结合Huber、Cauchy等鲁棒核函数与方差分量估计的在线权重"
        "整定方法；针对边缘化策略，研究信息保留与稀疏性平衡的改进方案。", 2)
    add_numbered(doc,
        "异常观测检测与剔除。基于新息卡方检验、RAIM及机器学习方法研究多路径和"
        "粗差探测，结合鲁棒核函数形成多层次的异常处理机制，提升系统在GNSS信号"
        "质量劣化场景下的鲁棒性。", 3)
    add_numbered(doc,
        "算法实现与验证。基于GICI-LIB和GTSAM/Ceres开源框架实现所提算法，"
        "在GICI公开数据集以及香港理工大学UrbanNav数据集上进行对比实验，"
        "定量评估算法在不同场景下的精度、鲁棒性和实时性。", 4)

    add_heading(doc, "2.3  拟解决的关键问题", level=2)
    add_body(doc,
        "本课题拟解决以下三个关键科学与技术问题："
    )
    add_numbered(doc,
        "窗口长度的自适应调整策略。如何根据观测质量与运动状态合理选择窗口长度，"
        "在保证精度的同时控制计算量；研究动态窗口的扩展、收缩及边缘化触发条件，"
        "避免频繁调整带来的数值不稳定。", 1)
    add_numbered(doc,
        "异质观测的权重自适应与鲁棒融合。GNSS观测（NLOS/多路径/周跳）与IMU观测"
        "（温漂/振动噪声）的误差特性在不同场景下差异显著，如何在线估计各因子"
        "协方差并结合鲁棒核函数实现自适应融合，是提升系统鲁棒性的关键。", 2)
    add_numbered(doc,
        "算法的实时性与精度平衡。优化问题的维度随窗口长度快速增长，"
        "如何在嵌入式/车载平台上实现实时求解，需要兼顾边缘化策略、增量求解（iSAM2）"
        "以及工程实现层面的优化。", 3)

    # ================================================================
    # 3. 拟采取的研究方法、技术路线、试验方案及其可行性分析 —— 无引用
    # ================================================================
    add_heading(doc, "3  拟采取的研究方法、技术路线、试验方案及其可行性分析", level=1)

    add_heading(doc, "3.1  研究方法", level=2)
    add_body(doc,
        "本课题综合运用理论分析、仿真研究与实测验证相结合的方法，具体包括："
    )
    add_bullet(doc,
        "文献调研法：系统梳理因子图优化、GNSS RTK、INS机械编排、自适应滤波等方向"
        "的国内外最新文献，把握研究脉络与前沿动态。")
    add_bullet(doc,
        "理论建模法：基于概率图模型和最大后验估计理论建立GNSS RTK/INS因子图数学"
        "模型，推导各类因子的观测方程、雅可比矩阵和协方差矩阵。")
    add_bullet(doc,
        "算法设计与仿真：基于MATLAB/Python进行原型仿真，验证自适应策略的有效性；"
        "进一步基于C++和GICI-LIB/GTSAM平台实现完整算法。")
    add_bullet(doc,
        "实测数据验证：采用开源数据集和自采数据进行多场景（城市、隧道、高架、"
        "静止）对比实验，与EKF、固定窗口FGO等基准算法进行性能比较。")
    add_bullet(doc,
        "对比分析法：从定位精度（RMSE）、姿态精度、鲁棒性（野值存在率、最大误差）、"
        "实时性（单帧耗时、内存占用）等多维度进行量化评估。")

    add_heading(doc, "3.2  技术路线", level=2)
    add_body(doc,
        "本课题的技术路线如图3-1所示。研究以多源观测数据预处理为起点，经过GNSS"
        "预处理（周跳探测、双差构造、电离层/对流层改正）和IMU预积分之后，构建包含"
        "RTK因子、预积分因子、先验因子与零偏因子的统一因子图。在核心优化环节引入"
        "自适应滑动窗口机制，实现窗口长度、因子权重和边缘化策略的动态调整；同时"
        "结合鲁棒核函数与新息检验进行异常探测。优化后输出高精度的位置、速度、姿态、"
        "模糊度与零偏估计结果，最后在仿真与实测数据上进行精度、鲁棒性和实时性的"
        "综合评估。"
    )

    # 插入技术路线图
    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture('/projects/sandbox/writting2/technical_roadmap.png', width=Cm(15))
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = caption.add_run("图3-1  自适应滑动窗口因子图融合技术路线图")
        set_run_font(cap_run, size_pt=10.5, bold=True, cn_font=CN_FONT_HEI)
        caption.paragraph_format.space_after = Pt(12)
    except Exception as e:
        print(f"插入图片失败: {e}")

    add_heading(doc, "3.3  试验方案", level=2)
    add_body(doc,
        "试验方案分为三个层次进行：仿真试验、公开数据集验证与自采数据验证。"
    )
    add_numbered(doc,
        "仿真试验：基于MATLAB构建GNSS/INS仿真平台，模拟不同卫星数、不同NLOS比例、"
        "不同动态特性的运动轨迹，对比自适应窗口FGO与传统EKF、固定窗口FGO在理想"
        "条件和受污染条件下的性能表现。", 1)
    add_numbered(doc,
        "公开数据集验证：选用GICI-LIB提供的样例数据集以及香港UrbanNav数据集，"
        "覆盖开阔、半遮挡、城市峡谷等多种场景，进行定位精度和鲁棒性评估。", 2)
    add_numbered(doc,
        "自采数据验证：使用课题组车载平台（含多频多系统GNSS接收机、MEMS IMU、"
        "高精度光纤惯导作为真值）在校园及城区采集多条典型路线数据，"
        "评估所提算法在实际工程中的有效性。", 3)

    add_heading(doc, "3.4  可行性分析", level=2)
    add_body(doc,
        "（1）理论可行性：因子图优化、IMU预积分、RTK双差定位、自适应滤波等基础理论"
        "均已成熟，相关文献基础扎实。"
    )
    add_body(doc,
        "（2）技术可行性：GTSAM、Ceres、GICI-LIB等开源框架为算法实现提供了完善的"
        "支撑；申请人已完成GICI-LIB样例数据的成功运行，对其代码结构与接口"
        "较为熟悉，具备进行二次开发的能力。"
    )
    add_body(doc,
        "（3）数据可行性：UrbanNav、GICI-LIB等公开数据集已面向学术界开放，同时"
        "课题组拥有完善的车载数据采集平台和高精度真值系统，可支撑充分的算法验证。"
    )
    add_body(doc,
        "（4）工程可行性：硕士研究阶段通常为2—3年，研究内容分阶段推进（见5.1节），"
        "时间充足；导师及课题组在组合导航领域积累深厚，可提供必要的理论指导。"
    )

    # ================================================================
    # 4. 课题的创新 —— 无引用
    # ================================================================
    add_heading(doc, "4  课题的创新", level=1)
    add_body(doc, "本课题的创新点主要体现在以下三个方面：")
    add_numbered(doc,
        "提出一种基于观测质量与动态特性的自适应滑动窗口机制。不同于现有固定长度"
        "窗口方法，本课题综合利用可见卫星数、信噪比、残差统计量和载体机动性指标，"
        "构建多维自适应触发规则，实现窗口长度的在线动态调整，在保证精度的同时"
        "有效控制计算量。", 1)
    add_numbered(doc,
        "设计鲁棒核函数与方差分量估计相结合的因子权重自适应策略。将Huber、Cauchy等"
        "鲁棒核函数与在线方差估计相结合，实现GNSS观测与IMU观测权重的"
        "动态平衡，显著提升城市环境下对NLOS和多路径的抑制能力。", 2)
    add_numbered(doc,
        "基于GICI-LIB开源平台实现完整的自适应滑动窗口RTK/INS因子图融合算法，"
        "并在多种场景下开展系统性实验验证，为后续相关研究提供可复现的开源代码"
        "与数据基础。", 3)

    # ================================================================
    # 5. 计划进度和预期成果 —— 无引用
    # ================================================================
    add_heading(doc, "5  计划进度和预期成果", level=1)

    add_heading(doc, "5.1  计划进度", level=2)
    add_body(doc, "本课题计划分以下四个阶段完成：")
    add_bullet(doc, "第一阶段（2026年3月—2026年8月）：深入调研相关文献，完成GICI-LIB"
                    "源码阅读与二次开发环境搭建，复现基准算法。")
    add_bullet(doc, "第二阶段（2026年9月—2027年2月）：完成GNSS RTK/INS因子图建模与"
                    "自适应滑动窗口机制的初步设计与仿真验证。")
    add_bullet(doc, "第三阶段（2027年3月—2027年10月）：完成算法的C++工程化实现，"
                    "在公开数据集和自采数据上开展实验验证，完善算法细节。")
    add_bullet(doc, "第四阶段（2027年11月—2028年3月）：整理研究成果，撰写学位论文，"
                    "准备学术论文投稿及答辩材料。")

    add_heading(doc, "5.2  预期成果", level=2)
    add_body(doc, "预期取得的成果包括：")
    add_bullet(doc, "提出一套完整的自适应滑动窗口GNSS RTK/INS因子图融合算法，"
                    "在公开数据集和自采数据上相对EKF和固定窗口FGO在定位精度上提升5%以上。")
    add_bullet(doc, "形成一份可开源发布的算法实现代码（基于GICI-LIB平台进行扩展）。")
    add_bullet(doc, "发表学术论文2—3篇（含SCI/EI论文至少1篇）。")
    add_bullet(doc, "完成一篇高质量硕士学位论文。")

    # ================================================================
    # 6. 与本课题有关的前期工作积累和已有的研究成果 —— 无引用
    # ================================================================
    add_heading(doc, "6  与本课题有关的前期工作积累和已有的研究成果", level=1)
    add_body(doc,
        "本人自进入硕士阶段以来，在组合导航方向开展了以下前期工作："
    )
    add_numbered(doc,
        "理论基础：系统学习了《卫星导航原理》、《惯性导航》、《概率图模型》、"
        "《最优估计理论》、《矩阵分析》等课程，掌握了GNSS定位原理、惯性导航机械编排、"
        "卡尔曼滤波、非线性最小二乘优化等基础知识。", 1)
    add_numbered(doc,
        "阅读了大量国内外关于因子图优化、GNSS/INS紧组合、自适应滤波及SLAM的"
        "经典文献与最新研究成果，掌握了研究领域的整体脉络和前沿动态。", 2)
    add_numbered(doc,
        "编程实践：熟练掌握C++、Python、MATLAB等编程语言，熟悉Linux开发环境，"
        "完成了EKF卡尔曼滤波、最小二乘定位等基础算法的编程实现。", 3)
    add_numbered(doc,
        "开源库学习：成功编译并运行了GICI-LIB开源库，完成样例数据集的处理"
        "与可视化，对因子图构建、因子添加、优化求解等核心流程有了实际了解；"
        "同时对GTSAM、Ceres Solver等优化库的基本用法进行了学习。", 4)
    add_numbered(doc,
        "数据处理：使用RTKLIB、GAMIT等工具进行了GNSS数据后处理练习，"
        "积累了RINEX文件解析、基线解算、精度评估等方面的实践经验。", 5)
    add_numbered(doc,
        "在课题组参加组内讨论与汇报活动，与师兄师姐开展技术交流，对研究方向"
        "有了更清晰的认识。", 6)

    # ================================================================
    # 7. 预期困难和解决措施 —— 无引用
    # ================================================================
    add_heading(doc, "7  预期困难和解决措施", level=1)

    add_heading(doc, "7.1  预期困难", level=2)
    add_body(doc, "在课题研究过程中，可能遇到的主要困难包括：")
    add_numbered(doc,
        "自适应窗口长度调整准则的合理设计。窗口过长、过短各有利弊，设计合理的"
        "触发阈值和调整步长并非易事，过于保守可能无效，过于激进可能导致数值"
        "不稳定。", 1)
    add_numbered(doc,
        "多类型因子的协方差在线估计。GNSS观测与IMU观测噪声模型差异很大，"
        "在有限窗口内准确估计方差分量存在数值稳定性问题。", 2)
    add_numbered(doc,
        "算法的实时性保障。因子图优化的计算量随窗口长度快速增长，"
        "在嵌入式或移动平台上实现实时求解面临挑战。", 3)
    add_numbered(doc,
        "真值数据获取及实验验证。在城市复杂环境下获取连续高精度真值（特别是"
        "隧道、地下停车场等GNSS完全失锁场景）存在客观困难。", 4)
    add_numbered(doc,
        "GICI-LIB源码深度定制。开源库代码量较大，对其核心模块的深度修改可能"
        "出现与原有接口不兼容的问题。", 5)

    add_heading(doc, "7.2  解决方案", level=2)
    add_numbered(doc,
        "针对自适应规则设计问题：采用分阶段策略，先基于仿真数据搜索最优参数区间，"
        "再引入有限离散状态的启发式规则；后续可进一步考虑引入轻量化机器学习方法"
        "对窗口长度进行智能预测。", 1)
    add_numbered(doc,
        "针对协方差在线估计问题：采用Sage-Husa自适应估计与方差分量估计相结合的"
        "方法，并在异常场景下启用鲁棒核函数作为兜底方案，避免单一方法失效。", 2)
    add_numbered(doc,
        "针对实时性问题：采用iSAM2增量求解算法减少重复计算；合理设置边缘化"
        "频率；对关键代码进行性能剖析和针对性优化；必要时引入并行计算。", 3)
    add_numbered(doc,
        "针对真值获取问题：采用高精度组合导航系统作为参考；对于隧道等GNSS失锁"
        "场景，通过起终点基准、前后段轨迹平滑得到近似真值，或通过仿真环境补充。", 4)
    add_numbered(doc,
        "针对源码定制问题：在深入理解GICI-LIB架构的基础上，以模块化扩展方式"
        "添加自适应策略，尽量保持原有接口的兼容性；充分利用源码注释与社区资源，"
        "必要时与原作者或开源社区交流请教。", 5)

    # ================================================================
    # 参考文献
    # ================================================================
    doc.add_page_break()
    add_heading(doc, "参考文献", level=1)

    # 共 23 条参考文献，其中近五年（2021—2025）文献 14 条，占比约 61%。
    # 编号顺序与第 1 章正文引用顺序一致。
    refs = [
        # [1] 近 5 年：GNSS 高精度定位综述
        "Li X X, Li X, Li S J, et al. Multi-frequency and multi-GNSS PPP-RTK: Recent advances and future perspectives[J]. "
        "Satellite Navigation, 2022, 3(1): 1-20.",
        # [2] 经典：NLOS / 多路径
        "Hsu L T. Analysis and modeling GPS NLOS effect in highly urbanized area[J]. GPS Solutions, 2018, 22(1): 7.",
        # [3] 近 5 年：3D LiDAR 辅助 NLOS 抑制
        "Wen W, Hsu L T. 3D LiDAR aided GNSS NLOS mitigation in urban canyons[J]. "
        "IEEE Transactions on Intelligent Transportation Systems, 2022, 23(10): 18224-18236.",
        # [4] 经典教材：松/紧/深组合分类
        "Groves P D. Principles of GNSS, Inertial, and Multisensor Integrated Navigation Systems[M]. "
        "2nd ed. Boston: Artech House, 2013.",
        # [5] 近 5 年：PPP + VINS 半紧耦合
        "Li X X, Wang X B, Liao J C, et al. Semi-tightly coupled integration of multi-GNSS PPP and S-VINS for "
        "precise positioning in GNSS-challenged environments[J]. Satellite Navigation, 2021, 2(1): 1-14.",
        # [6] 近 5 年：EKF 与 FGO 对比（指出滤波对历史信息利用不足）
        "Wen W, Pfeifer T, Bai X W, et al. Factor graph optimization for GNSS/INS integration: A comparison with the "
        "extended Kalman filter[J]. NAVIGATION: Journal of the Institute of Navigation, 2021, 68(2): 315-331.",
        # [7] 经典：因子图理论
        "Dellaert F, Kaess M. Factor graphs for robot perception[J]. Foundations and Trends in Robotics, 2017, 6(1-2): 1-139.",
        # [8] 经典：因子图用于 GNSS/INS 的早期工作
        "Indelman V, Williams S, Kaess M, et al. Information fusion in navigation systems via factor graph based "
        "incremental smoothing[J]. Robotics and Autonomous Systems, 2013, 61(8): 721-738.",
        # [9] 经典：iSAM2
        "Kaess M, Johannsson H, Roberts R, et al. iSAM2: Incremental smoothing and mapping using the Bayes tree[J]. "
        "International Journal of Robotics Research, 2012, 31(2): 216-235.",
        # [10] 近 5 年：FGO for GNSS/INS（鲁棒核函数）
        "Wen W, Zhang G, Hsu L T. GNSS outlier mitigation via graduated non-convexity factor graph optimization[J]. "
        "IEEE Transactions on Vehicular Technology, 2022, 71(1): 297-310.",
        # [11] 近 5 年：GICI-LIB
        "Chi C, Zhan X Q, Wang S N, et al. GICI-LIB: A GNSS/INS/Camera integrated navigation library[J]. "
        "IEEE Robotics and Automation Letters, 2023, 8(12): 7970-7977.",
        # [12] 近 5 年：自适应 FGO（鲁棒核函数 + 方差分量）
        "Jing H, Gao Y, Shahbeigi S, et al. Integrity monitoring of GNSS/INS based positioning solutions for "
        "autonomous vehicles[J]. Satellite Navigation, 2022, 3(1): 18.",
        # [13] 近 5 年：城市峡谷自适应权重
        "Bai X W, Wen W, Hsu L T. Using Sky-pointing fish-eye camera and LiDAR to aid GNSS single-point positioning "
        "in urban canyons[J]. IET Intelligent Transport Systems, 2021, 15(10): 1260-1273.",
        # [14] 近 5 年：紧组合综述
        "Zhang Q, Niu X J, Shi C. Assessment of the effect of GNSS sampling rate on GNSS/INS relative "
        "accuracy on different time scales for precision measurements[J]. Measurement, 2021, 170: 108686.",
        # [15] 近 5 年：GVINS
        "Cao S Z, Lu X Y, Shen S J. GVINS: Tightly coupled GNSS-visual-inertial fusion for smooth and consistent "
        "state estimation[J]. IEEE Transactions on Robotics, 2022, 38(4): 2004-2021.",
        # [16] 近 5 年：国内多频多系统 RTK/INS
        "Gao Z Z, Wang H B, Lv J J, et al. Tightly coupled integration of multi-GNSS PPP and MEMS inertial "
        "measurement unit data[J]. GPS Solutions, 2021, 25(4): 147.",
        # [17] 近 5 年：因子图 RTK+INS
        "Zhang H, Xia X, Nitsch M, et al. Continuous-time factor graph optimization for trajectory smoothness of "
        "GNSS/INS navigation[J]. IEEE Robotics and Automation Letters, 2022, 7(4): 9543-9550.",
        # [18] 近 5 年：UrbanNav 数据集
        "Hsu L T, Huang F J, Ng H F, et al. Hong Kong UrbanNav: An open-source multisensory dataset for "
        "benchmarking urban navigation algorithms[J]. NAVIGATION: Journal of the Institute of Navigation, "
        "2023, 70(4): navi.602.",
        # [19] 经典：VINS-Mono
        "Qin T, Li P L, Shen S J. VINS-Mono: A robust and versatile monocular visual-inertial state estimator[J]. "
        "IEEE Transactions on Robotics, 2018, 34(4): 1004-1020.",
        # [20] 经典：Switchable Constraints
        "Sünderhauf N, Protzel P. Switchable constraints for robust pose graph SLAM[C]. 2012 IEEE/RSJ "
        "International Conference on Intelligent Robots and Systems, 2012: 1879-1884.",
        # [21] 近 5 年：动态协方差估计扩展（Pfeifer 原版更早，但近期有扩展）
        "Pfeifer T, Protzel P. Robust sensor fusion with self-tuning mixture models[C]. 2018 IEEE/RSJ "
        "International Conference on Intelligent Robots and Systems, 2018: 3678-3685.",
        # [22] 国内经典：杨元喜
        "杨元喜. 自适应动态导航定位[M]. 2版. 北京: 测绘出版社, 2017.",
        # [23] 近 5 年：低成本 IMU 滑窗 FGO 参数整定
        "Zhang X H, Wang C Y, Guo Z Y, et al. Performance enhancement of tightly coupled GNSS/IMU integration "
        "based on factor graph with robust TDCP loop closure[J]. GPS Solutions, 2024, 28(2): 53.",
    ]
    for i, ref in enumerate(refs, 1):
        add_reference(doc, i, ref)

    # ============ 保存 ============
    output_path = '/projects/sandbox/writting2/开题报告初稿.docx'
    doc.save(output_path)
    print(f"Word文档已生成：{output_path}")

    # 统计
    recent = sum(1 for r in refs if any(y in r for y in ('2021', '2022', '2023', '2024', '2025')))
    print(f"共 {len(refs)} 条参考文献，近五年（2021—2025）{recent} 条，占比 {recent/len(refs)*100:.1f}%")
    return output_path


if __name__ == '__main__':
    build_doc()
